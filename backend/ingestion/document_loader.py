"""
Document Loader - Load and parse various document types
"""

from typing import List, Dict, Any
import os
from pathlib import Path


class DocumentLoader:
    """
    Load documents from various file types
    """
    
    def __init__(self):
        self.supported_extensions = [".txt", ".md", ".pdf", ".docx", ".html"]
    
    async def load_from_file(self, file_path: str, file_content: bytes = None) -> Dict[str, Any]:
        """
        Load document from file
        
        Args:
            file_path: Path to file
            file_content: Optional file content bytes
            
        Returns:
            Dict with content and metadata
        """
        extension = Path(file_path).suffix.lower()
        
        if extension == ".txt" or extension == ".md":
            return await self._load_text(file_path, file_content)
        elif extension == ".pdf":
            return await self._load_pdf(file_path, file_content)
        elif extension == ".docx":
            return await self._load_docx(file_path, file_content)
        elif extension == ".html":
            return await self._load_html(file_path, file_content)
        else:
            # Try to load as text
            return await self._load_text(file_path, file_content)
    
    async def _load_text(self, file_path: str, file_content: bytes = None) -> Dict[str, Any]:
        """Load plain text or markdown"""
        try:
            if file_content:
                content = file_content.decode('utf-8')
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            return {
                "content": content,
                "metadata": {
                    "source": Path(file_path).name,
                    "type": "text"
                }
            }
        except Exception as e:
            raise ValueError(f"Error loading text file: {e}")
    
    async def _load_pdf(self, file_path: str, file_content: bytes = None) -> Dict[str, Any]:
        """Load PDF document"""
        try:
            import PyPDF2
            from io import BytesIO
            
            if file_content:
                pdf_file = BytesIO(file_content)
            else:
                pdf_file = open(file_path, 'rb')
            
            reader = PyPDF2.PdfReader(pdf_file)
            content = ""
            
            for page_num, page in enumerate(reader.pages):
                content += f"\n--- Page {page_num + 1} ---\n"
                content += page.extract_text()
            
            if not file_content:
                pdf_file.close()
            
            return {
                "content": content,
                "metadata": {
                    "source": Path(file_path).name,
                    "type": "pdf",
                    "pages": len(reader.pages)
                }
            }
        except ImportError:
            return {
                "content": "PDF parsing not available. Install PyPDF2.",
                "metadata": {"source": Path(file_path).name, "type": "pdf"}
            }
        except Exception as e:
            raise ValueError(f"Error loading PDF: {e}")
    
    async def _load_docx(self, file_path: str, file_content: bytes = None) -> Dict[str, Any]:
        """Load Word document"""
        try:
            import docx
            from io import BytesIO
            
            if file_content:
                doc = docx.Document(BytesIO(file_content))
            else:
                doc = docx.Document(file_path)
            
            content = "\n".join([para.text for para in doc.paragraphs])
            
            return {
                "content": content,
                "metadata": {
                    "source": Path(file_path).name,
                    "type": "docx"
                }
            }
        except ImportError:
            return {
                "content": "DOCX parsing not available. Install python-docx.",
                "metadata": {"source": Path(file_path).name, "type": "docx"}
            }
        except Exception as e:
            raise ValueError(f"Error loading DOCX: {e}")
    
    async def _load_html(self, file_path: str, file_content: bytes = None) -> Dict[str, Any]:
        """Load HTML document"""
        try:
            from bs4 import BeautifulSoup
            
            if file_content:
                html = file_content.decode('utf-8')
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    html = f.read()
            
            soup = BeautifulSoup(html, 'html.parser')
            content = soup.get_text(separator='\n', strip=True)
            
            return {
                "content": content,
                "metadata": {
                    "source": Path(file_path).name,
                    "type": "html"
                }
            }
        except ImportError:
            return {
                "content": "HTML parsing not available. Install beautifulsoup4.",
                "metadata": {"source": Path(file_path).name, "type": "html"}
            }
        except Exception as e:
            raise ValueError(f"Error loading HTML: {e}")
